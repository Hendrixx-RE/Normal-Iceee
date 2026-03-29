"""
Multi-format file text extraction.

Supported types:
  PDF   → existing OCR pipeline (PyMuPDF native text or doctr for scanned pages)
  Image → doctr OCR → text (same engine used for scanned PDF pages)
           falls back to returning image for Gemini multimodal if quality is poor
  Word  → python-docx paragraph + table extraction
  Excel → openpyxl sheet/row extraction
  CSV   → built-in csv module
"""
import csv
import io
import logging
from pathlib import Path
from typing import List, Optional, Tuple

from PIL import Image

from app.services.ocr_strategies.image_based import DocTROCREngine
from app.services.ocr_strategies.quality_checker import TextQualityChecker

logger = logging.getLogger(__name__)

# If doctr quality is below this, also return the image for Gemini multimodal fallback
DOCTR_QUALITY_THRESHOLD = 40.0

SUPPORTED_EXTENSIONS = {
    '.pdf', '.jpg', '.jpeg', '.png', '.webp', '.tiff', '.tif', '.bmp',
    '.docx', '.xlsx', '.xls', '.csv',
}

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.webp', '.tiff', '.tif', '.bmp'}


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == '.pdf':
        return 'pdf'
    if ext in IMAGE_EXTENSIONS:
        return 'image'
    if ext == '.docx':
        return 'docx'
    if ext in {'.xlsx', '.xls'}:
        return 'excel'
    if ext == '.csv':
        return 'csv'
    raise ValueError(f"Unsupported file type: '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# Image
# ---------------------------------------------------------------------------

def extract_from_image(file_bytes: bytes) -> Tuple[str, Optional[List[Image.Image]]]:
    """
    Extract text from an image file using doctr — the same engine used for
    scanned PDF pages.

    If doctr quality is poor (< DOCTR_QUALITY_THRESHOLD), also returns the
    image so the caller can pass it to Gemini multimodal as a fallback.
    """
    img = Image.open(io.BytesIO(file_bytes)).convert('RGB')
    logger.info(f"Image loaded: {img.size[0]}x{img.size[1]} px")

    text = ''
    try:
        text = DocTROCREngine.extract_text(img)
        logger.info(f"doctr image OCR: {len(text)} chars")
    except Exception as exc:
        logger.error(f"doctr failed on image: {exc} — will use Gemini multimodal")

    quality = TextQualityChecker.get_quality_score(text) if text else 0.0
    logger.info(f"Image OCR quality: {quality:.1f}/100")

    # Return the image alongside text if quality is too low for text-only Gemini
    fallback_images = [img] if quality < DOCTR_QUALITY_THRESHOLD else None
    return text, fallback_images


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------

def extract_from_docx(file_bytes: bytes) -> str:
    """Extract paragraphs and table cells from a Word document."""
    from docx import Document  # python-docx — imported lazily to avoid startup cost
    doc = Document(io.BytesIO(file_bytes))
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(' | '.join(cells))

    text = '\n'.join(parts)
    logger.info(f"DOCX extraction: {len(text)} chars from {len(doc.paragraphs)} paragraphs")
    return text


# ---------------------------------------------------------------------------
# Excel (.xlsx / .xls)
# ---------------------------------------------------------------------------

def extract_from_excel(file_bytes: bytes, filename: str = '') -> str:
    """Convert all sheets to labelled plain-text rows. Supports .xlsx and .xls."""
    ext = Path(filename).suffix.lower() if filename else '.xlsx'

    if ext == '.xls':
        # xlrd handles the legacy .xls binary format; openpyxl cannot
        import xlrd
        wb_xls = xlrd.open_workbook(file_contents=file_bytes)
        parts: List[str] = []
        for sheet in wb_xls.sheets():
            parts.append(f"=== Sheet: {sheet.name} ===")
            for rx in range(sheet.nrows):
                cells = [str(sheet.cell_value(rx, cx)).strip()
                         for cx in range(sheet.ncols)
                         if str(sheet.cell_value(rx, cx)).strip()]
                if cells:
                    parts.append(' | '.join(cells))
        logger.info(f"XLS extraction: {wb_xls.nsheets} sheet(s)")
        return '\n'.join(parts)

    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts: List[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if cells:
                parts.append(' | '.join(cells))
                row_count += 1
        logger.info(f"Sheet '{sheet_name}': {row_count} non-empty rows")

    wb.close()
    return '\n'.join(parts)


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------

def extract_from_csv(file_bytes: bytes) -> str:
    """Convert CSV rows to pipe-delimited plain text."""
    text = file_bytes.decode('utf-8', errors='replace')
    reader = csv.reader(io.StringIO(text))
    rows = [
        ' | '.join(cell.strip() for cell in row)
        for row in reader
        if any(cell.strip() for cell in row)
    ]
    logger.info(f"CSV extraction: {len(rows)} rows")
    return '\n'.join(rows)


# ---------------------------------------------------------------------------
# Unified entry point (for non-PDF types)
# ---------------------------------------------------------------------------

def extract_non_pdf(file_bytes: bytes, filename: str) -> Tuple[str, Optional[List[Image.Image]]]:
    """
    Extract text (and optionally images) from non-PDF files.

    Returns:
        (text, images) — images is non-None only for image files
    """
    file_type = get_file_type(filename)

    if file_type == 'image':
        text, images = extract_from_image(file_bytes)
        return text, images

    if file_type == 'docx':
        return extract_from_docx(file_bytes), None

    if file_type == 'excel':
        return extract_from_excel(file_bytes, filename), None

    if file_type == 'csv':
        return extract_from_csv(file_bytes), None

    raise ValueError(f"Unexpected file type for non-PDF extractor: {file_type}")
